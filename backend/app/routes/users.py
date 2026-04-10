from flask import Blueprint, request, jsonify, url_for, send_file, make_response
from flask_jwt_extended import jwt_required, get_jwt_identity
from app import db
from app.models import User
from app.utils.scoping import scope_query, get_current_workspace_id
from werkzeug.utils import secure_filename
from datetime import datetime
import secrets
import os
import mimetypes
import io
import json
from app.utils.audit import log_audit_event
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

users_bp = Blueprint('users', __name__)

def is_admin(user):
    """Check if user is admin"""
    return user and user.role in ['admin', 'super_admin']

# Route removed: handled by user_create_bp in routes/user_create.py
# Route removed: handled by user_create_bp in routes/user_create.py




@users_bp.route('', methods=['GET'])
@jwt_required()
def get_users():
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user:
        return jsonify({'error': 'User not found'}), 404
    
    # Only admins can list users
    if not is_admin(current_user):
        return jsonify({'error': 'Unauthorized'}), 403
    
    if current_user.role == 'super_admin':
        workspace_id = request.args.get('workspace_id')
        if workspace_id:
             users = User.query.filter_by(workspace_id=workspace_id).all()
        else:
             users = User.query.all()
    else:
        # Regular admin - use utility for automatic scoping
        users = scope_query(User.query, User).filter(User.role != 'super_admin').all()
        
    return jsonify([user.to_dict() for user in users]), 200

@users_bp.route('/<int:user_id>', methods=['GET'])
@jwt_required()
def get_user(user_id):
    current_user_id = get_jwt_identity()
    user = scope_query(User.query, User).filter_by(id=user_id).first()
    
    if not user:
        return jsonify({'error': 'User not found or access denied'}), 404
    
    return jsonify(user.to_dict()), 200

@users_bp.route('/me', methods=['PUT'])
@jwt_required()
def update_current_user():
    current_user_id = get_jwt_identity()
    user = User.query.get(current_user_id)
    
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    data = request.get_json()
    
    if 'first_name' in data:
        user.first_name = data['first_name']
    if 'last_name' in data:
        user.last_name = data['last_name']
    if 'avatar_url' in data:
        user.avatar_url = data['avatar_url']
    
    db.session.commit()
    
    return jsonify(user.to_dict()), 200

@users_bp.route('/me/avatar', methods=['POST'])
@jwt_required()
def upload_avatar():
    """Upload profile picture for current user"""
    current_user_id = get_jwt_identity()
    user = User.query.get(current_user_id)
    
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Check if it's an image
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'svg'}
    if '.' not in file.filename:
        return jsonify({'error': 'Invalid file type'}), 400
    
    ext = file.filename.rsplit('.', 1)[1].lower()
    if ext not in allowed_extensions:
        return jsonify({'error': 'Only image files are allowed (png, jpg, jpeg, gif, webp, svg)'}), 400
    
    # Generate unique filename
    filename = secure_filename(file.filename)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
    unique_filename = f"{timestamp}{current_user_id}_{filename}"
    
    # Storage Logic
    file_url = None
    from app.utils.storage import upload_fileobj, get_public_url
    from flask import current_app
    
    # Cloud Storage Path
    if current_app.config.get('S3_ENDPOINT') and current_app.config.get('S3_BUCKET'):
        # Upload to S3
        file.seek(0)
        key = f"avatars/{unique_filename}"
        if upload_fileobj(file, key):
            file_url = get_public_url(key)
            
    # Fallback to local
    if not file_url:
        # Create avatars directory if it doesn't exist
        current_file_dir = os.path.dirname(os.path.abspath(__file__))
        backend_dir = os.path.dirname(os.path.dirname(current_file_dir))
        avatars_folder = os.path.join(backend_dir, 'uploads', 'avatars')
        os.makedirs(avatars_folder, exist_ok=True)
        
        file_path = os.path.join(avatars_folder, unique_filename)
        file.seek(0)
        file.save(file_path)
        file_url = f"/api/files/avatar/{unique_filename}"
    
    # Update user's avatar URL
    user.avatar_url = file_url
    db.session.commit()
    
    return jsonify({
        'message': 'Avatar uploaded successfully',
        'avatar_url': file_url,
        'user': user.to_dict()
    }), 200

# Route removed and moved to top for reliability

@users_bp.route('/search', methods=['GET'])
@jwt_required()
def search_users():
    query = request.args.get('q', '')
    global_search = request.args.get('global', 'false').lower() == 'true'
    
    if not query or len(query) < 2:
        return jsonify({'error': 'Query too short'}), 400
    
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    # 1. First, search within the current workspace (priority)
    workspace_query = User.query.filter(
        User.workspace_id == current_user.workspace_id,
        (
            (User.username.ilike(f'%{query}%')) |
            (User.email.ilike(f'%{query}%')) |
            (User.first_name.ilike(f'%{query}%')) |
            (User.last_name.ilike(f'%{query}%'))
        )
    )
    
    if current_user.role != 'super_admin':
        workspace_query = workspace_query.filter(User.role != 'super_admin')
        
    workspace_users = workspace_query.limit(20).all()
    
    # 2. If global search enabled or fewer than 5 users found, search across other workspaces
    # (Allow cross-workspace search for common identity fields)
    other_users = []
    if (global_search or len(workspace_users) < 3) and len(query) >= 3:
        # Cross-workspace: allow partial name/email matches but limit results for privacy
        other_query = User.query.filter(
            User.workspace_id != current_user.workspace_id,
            (
                (User.email.ilike(f'%{query}%')) | 
                (User.username.ilike(f'%{query}%')) |
                (User.first_name.ilike(f'%{query}%')) |
                (User.last_name.ilike(f'%{query}%'))
            )
        )
        # Exclude super admins from general cross-workspace discovery
        other_query = other_query.filter(User.role != 'super_admin')
        other_users = other_query.limit(10).all()
    
    # Combined results, workspace users first
    seen_ids = set(u.id for u in workspace_users)
    combined_users = workspace_users + [u for u in other_users if u.id not in seen_ids]
    
    return jsonify([user.to_dict() for user in combined_users]), 200

    return jsonify([user.to_dict() for user in combined_users]), 200


@users_bp.route('/list-teachers', methods=['GET'])
@jwt_required()
def get_public_teachers():
    """Get list of teachers (public for auth users)"""
    teachers = User.query.filter_by(role='teacher', is_active=True).all()
    # Return minimal info
    return jsonify([{
        'id': u.id,
        'first_name': u.first_name,
        'last_name': u.last_name,
        'username': u.username,
        'role': u.role,
        'avatar_url': u.avatar_url
    } for u in teachers]), 200

@users_bp.route('/list-admins', methods=['GET'])
@jwt_required()
def get_public_admins():
    """Get list of admins (public for auth users)"""
    admins = User.query.filter_by(role='admin', is_active=True).all()
    return jsonify([{
        'id': u.id,
        'first_name': u.first_name,
        'last_name': u.last_name,
        'username': u.username,
        'role': u.role,
        'avatar_url': u.avatar_url
    } for u in admins]), 200

@users_bp.route('/students', methods=['GET'])
@jwt_required()
def get_students():
    """Get all students (admin only)"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    if current_user.role == 'super_admin':
        workspace_id = request.args.get('workspace_id')
        if workspace_id:
             students = User.query.filter_by(role='student', workspace_id=workspace_id).all()
        else:
             students = User.query.filter_by(role='student').all()
    else:
        students = scope_query(User.query, User).filter_by(role='student').all()
    return jsonify([user.to_dict() for user in students]), 200

@users_bp.route('/students', methods=['POST'])
@jwt_required()
def add_student():
    """Add a new student (admin only)"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    data = request.get_json()
    
    if not data.get('email'):
        return jsonify({'error': 'Email is required'}), 400
    
    # Check if user already exists
    existing_user = User.query.filter_by(email=data['email']).first()
    if existing_user:
        return jsonify({'error': 'User with this email already exists'}), 400
    
    # Generate username from email if not provided
    username = data.get('username') or data['email'].split('@')[0]
    
    # Ensure username is unique
    base_username = username
    counter = 1
    while User.query.filter_by(username=username).first():
        username = f"{base_username}{counter}"
        counter += 1
    
    student = User(
        username=username,
        email=data['email'],
        first_name=data.get('first_name', ''),
        last_name=data.get('last_name', ''),
        role='student',
        workspace_id=current_user.workspace_id,
        admin_id=current_user.id
    )
    
    # Set password if provided, otherwise generate a temporary one
    password = data.get('password') or secrets.token_urlsafe(12)
    student.set_password(password)
    
    db.session.add(student)
    db.session.commit()
    
    return jsonify({
        'message': 'Student added successfully',
        'user': student.to_dict(),
        'temporary_password': password if not data.get('password') else None
    }), 201

@users_bp.route('/students/<int:student_id>', methods=['DELETE'])
@jwt_required()
def remove_student(student_id):
    """Remove a student (admin only)"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    student = scope_query(User.query, User).filter_by(id=student_id).first()
    
    if not student:
        return jsonify({'error': 'Student not found or access denied'}), 404
    
    if student.role != 'student':
        return jsonify({'error': 'User is not a student'}), 400
    
    if student.id == current_user_id:
        return jsonify({'error': 'Cannot delete yourself'}), 400
    
    db.session.delete(student)
    db.session.commit()
    
    # Log audit event
    log_audit_event(
        user_id=current_user_id,
        action='delete_user',
        resource_type='user',
        resource_id=student_id,
        details={
            'deleted_user_email': student.email,
            'deleted_user_role': student.role,
            'method': 'remove_student'
        }
    )
    
    return jsonify({'message': 'Student removed successfully'}), 200

@users_bp.route('/students/<int:student_id>', methods=['PUT'])
@jwt_required()
def update_student(student_id):
    """Update student information (admin only)"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    student = User.query.get(student_id)
    
    if not student:
        return jsonify({'error': 'Student not found'}), 404
    
    if student.role != 'student':
        return jsonify({'error': 'User is not a student'}), 400
    
    data = request.get_json()
    
    if 'first_name' in data:
        student.first_name = data['first_name']
    if 'last_name' in data:
        student.last_name = data['last_name']
    if 'email' in data:
        # Prevent email changes as per institutional policy - login email is permanent
        pass
    if 'username' in data:
        # Check if username is already taken
        existing = User.query.filter_by(username=data['username']).first()
        if existing and existing.id != student_id:
            return jsonify({'error': 'Username already in use'}), 400
        student.username = data['username']
    if 'password' in data:
        student.set_password(data['password'])
    if 'is_active' in data:
        student.is_active = data['is_active']
    
    if 'workspace_id' in data:
        # Immutable Workspace Assignment: Only Super Admin can change workspace
        if current_user.role != 'super_admin':
            return jsonify({'error': 'Workspace assignment is immutable for regular admins'}), 403
            
        # Check if workspace exists
        if data['workspace_id']:
            from app.models import Workspace
            ws = Workspace.query.get(data['workspace_id'])
            if not ws:
                return jsonify({'error': 'Workspace not found'}), 404
            student.workspace_id = ws.id
        else:
            return jsonify({'error': 'Workspace cannot be None'}), 400
    
    db.session.commit()
    
    return jsonify({
        'message': 'Student updated successfully',
        'user': student.to_dict()
    }), 200

@users_bp.route('/<int:user_id>/role', methods=['PUT'])
@jwt_required()
def update_user_role(user_id):
    """Update user role - Admin only"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    user = User.query.get(user_id)
    
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    # Cannot change own role
    if user.id == current_user_id:
        return jsonify({'error': 'Cannot change your own role'}), 400
    
    data = request.get_json()
    new_role = data.get('role')
    
    if not new_role:
        return jsonify({'error': 'Role is required'}), 400
    
    if new_role not in ['admin', 'teacher', 'student']:
        return jsonify({'error': 'Invalid role. Must be admin, teacher, or student'}), 400
    
    user.role = new_role
    db.session.commit()
    
    # Emit Socket.IO event to notify the user of role change
    try:
        from app import socketio
        socketio.emit('user_role_updated', {
            'user_id': user_id,
            'new_role': new_role,
            'user': user.to_dict()
        }, room=f'user_{user_id}')
    except Exception as e:
        print(f"Error emitting role update event: {e}")
    
    return jsonify({
        'message': 'User role updated successfully',
        'user': user.to_dict()
    }), 200

@users_bp.route('/<int:user_id>/privileges', methods=['PUT'])
@jwt_required()
def update_user_privileges(user_id):
    """Update user privileges - Admin only, takes effect immediately"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    user = User.query.get(user_id)
    
    if not user:
        return jsonify({'error': 'User not found'}), 404

    # Security: Regular admins can only manage users in their workspace - simplified using utility
    if not scope_query(User.query, User).filter_by(id=user_id).first():
        return jsonify({'error': 'Unauthorized. User not found in your workspace'}), 403
    
    data = request.get_json()
    privileges = data.get('privileges', {})
    
    if not isinstance(privileges, dict):
        return jsonify({'error': 'Privileges must be a dictionary/object'}), 400
    
    # Validate privilege keys (list of common privileges)
    valid_privileges = [
        'can_manage_users',
        'can_manage_students',
        'can_manage_teachers',
        'can_manage_classes',
        'can_manage_groups',
        'can_manage_rooms',
        'can_manage_channels',
        'can_manage_system_settings',
        'can_export_data',
        'can_view_analytics',
        'can_manage_permissions',
        'can_broadcast_messages'
    ]
    
    # Filter to only valid privileges and ensure boolean values
    filtered_privileges = {}
    for key, value in privileges.items():
        if key in valid_privileges:
            filtered_privileges[key] = bool(value)
    
    user.privileges = filtered_privileges
    db.session.commit()
    
    # Emit Socket.IO event to notify the user of privilege change
    try:
        from app import socketio
        socketio.emit('user_privileges_updated', {
            'user_id': user_id,
            'new_privileges': filtered_privileges,
            'user': user.to_dict()
        }, room=f'user_{user_id}')
    except Exception as e:
        print(f"Error emitting privileges update event: {e}")
    
    return jsonify({
        'message': 'Privileges updated successfully',
        'user': user.to_dict()
    }), 200

@users_bp.route('/privileges/list', methods=['GET'])
@jwt_required()
def get_available_privileges():
    """Get list of available privileges - Admin only"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    privileges_list = [
        {
            'key': 'can_manage_users',
            'label': 'Manage Users',
            'description': 'Can create, edit, and delete users'
        },
        {
            'key': 'can_manage_students',
            'label': 'Manage Students',
            'description': 'Can manage student accounts and data'
        },
        {
            'key': 'can_manage_teachers',
            'label': 'Manage Teachers',
            'description': 'Can manage teacher accounts and assignments'
        },
        {
            'key': 'can_manage_classes',
            'label': 'Manage Classes',
            'description': 'Can create, edit, and delete classes/courses'
        },
        {
            'key': 'can_manage_groups',
            'label': 'Manage Groups',
            'description': 'Can create, edit, and delete groups'
        },
        {
            'key': 'can_manage_rooms',
            'label': 'Manage Video Conferences',
            'description': 'Can manage video conference rooms'
        },
        {
            'key': 'can_manage_channels',
            'label': 'Manage Chatrooms',
            'description': 'Can manage chat channels (but cannot access unless member)'
        },
        {
            'key': 'can_manage_system_settings',
            'label': 'Manage System Settings',
            'description': 'Can modify system-wide settings'
        },
        {
            'key': 'can_export_data',
            'label': 'Export Data',
            'description': 'Can export user and system data'
        },
        {
            'key': 'can_view_analytics',
            'label': 'View Analytics',
            'description': 'Can access analytics and reports'
        },
        {
            'key': 'can_manage_permissions',
            'label': 'Manage Permissions',
            'description': 'Can assign privileges to other users'
        },
        {
            'key': 'can_broadcast_messages',
            'label': 'Broadcast Messages',
            'description': 'Can send broadcast messages to all users'
        }
    ]
    
    return jsonify(privileges_list), 200

@users_bp.route('/export/<role>/<format>', methods=['GET'])
@jwt_required()
def export_users(role, format):
    """Export users (students or teachers) in Excel, PDF, or Word format (admin only)"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
    
    if role not in ['students', 'teachers']:
        return jsonify({'error': 'Invalid role. Use "students" or "teachers"'}), 400
    
    if format not in ['xlsx', 'pdf', 'docx']:
        return jsonify({'error': 'Invalid format. Use "xlsx", "pdf", or "docx"'}), 400
    
    # Get users based on role
    # Get users based on role
    role_filter = 'student' if role == 'students' else 'teacher'
    query = User.query.filter_by(role=role_filter)
    
    if current_user.role != 'super_admin':
        query = query.filter_by(workspace_id=current_user.workspace_id)
        
    users = query.all()
    
    if format == 'xlsx':
        return export_to_excel(users, role)
    elif format == 'pdf':
        return export_to_pdf(users, role)
    elif format == 'docx':
        return export_to_word(users, role)

def export_to_excel(users, role_name):
    """Export users to Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = role_name.title()
    
    # Headers
    headers = ['ID', 'Username', 'Email', 'First Name', 'Last Name', 'Role', 'Status', 'Created At']
    ws.append(headers)
    
    # Style header row
    from openpyxl.styles import Font, PatternFill
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
    
    # Add data
    for user in users:
        ws.append([
            user.id,
            user.username or '',
            user.email,
            user.first_name or '',
            user.last_name or '',
            user.role,
            'Active' if user.is_active else 'Inactive',
            user.created_at.strftime('%Y-%m-%d %H:%M:%S') if user.created_at else ''
        ])
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    filename = f"{role_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )

def export_to_pdf(users, role_name):
    """Export users to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    
    styles = getSampleStyleSheet()
    title = Paragraph(f"{role_name.title()} List", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))
    
    # Table data
    data = [['ID', 'Username', 'Email', 'First Name', 'Last Name', 'Role', 'Status']]
    
    for user in users:
        data.append([
            str(user.id),
            user.username or '',
            user.email,
            user.first_name or '',
            user.last_name or '',
            user.role,
            'Active' if user.is_active else 'Inactive'
        ])
    
    # Create table
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    
    filename = f"{role_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename
    )

def export_to_word(users, role_name):
    """Export users to Word format"""
    doc = Document()
    doc.add_heading(f'{role_name.title()} List', 0)
    
    # Create table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'Username'
    header_cells[2].text = 'Email'
    header_cells[3].text = 'First Name'
    header_cells[4].text = 'Last Name'
    header_cells[5].text = 'Role'
    header_cells[6].text = 'Status'
    
    # Add data rows
    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = str(user.id)
        row_cells[1].text = user.username or ''
        row_cells[2].text = user.email
        row_cells[3].text = user.first_name or ''
        row_cells[4].text = user.last_name or ''
        row_cells[5].text = user.role
        row_cells[6].text = 'Active' if user.is_active else 'Inactive'
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{role_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@users_bp.route('/active', methods=['GET'])
@jwt_required()
def get_active_users():
    '''Get active users (online/recent) - Admin only'''
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized'}), 403
        
    query = User.query.filter_by(is_active=True)
    
    if current_user.role == 'super_admin':
        workspace_id = request.args.get('workspace_id')
        if workspace_id:
            query = query.filter_by(workspace_id=workspace_id)
    else:
        # Regular admin - use utility
        query = scope_query(query, User).filter(User.role != 'super_admin')
        
    users = query.all()
    return jsonify([user.to_dict() for user in users]), 200


@users_bp.route('/<int:user_id>', methods=['DELETE'])
@jwt_required()
def delete_user(user_id):
    """Delete a user account (Admin/Super Admin only)"""
    current_user_id = get_jwt_identity()
    current_user = User.query.get(current_user_id)
    
    if not current_user or not is_admin(current_user):
        return jsonify({'error': 'Unauthorized. Admin access required'}), 403
        
    user_to_delete = User.query.get(user_id)
    if not user_to_delete:
        return jsonify({'error': 'User not found'}), 404
        
    # Permission checks
    is_super = current_user.role == 'super_admin' or getattr(current_user, 'platform_role', None) == 'SUPER_ADMIN'
    
    if not is_super:
        # Check if they are the primary admin (Recognized Head) of their workspace
        from app.models import Workspace
        workspace = Workspace.query.get(current_user.workspace_id) if current_user.workspace_id else None
        is_workspace_head = workspace and workspace.admin_id == current_user.id
        
        # Regular administrative activities should only be carried out by the recognized head
        if not is_workspace_head:
            return jsonify({'error': 'Unauthorized. Only the recognized Workspace Admin or a Super Admin can delete accounts.'}), 403

        # Recognized Workspace Admin can ONLY delete users in their own workspace
        if user_to_delete.workspace_id != current_user.workspace_id:
            return jsonify({'error': 'Unauthorized. Cannot delete user from another workspace.'}), 403
            
        # Recognized Workspace Admin CANNOT delete super admins
        if user_to_delete.role == 'super_admin' or getattr(user_to_delete, 'platform_role', None) == 'SUPER_ADMIN':
            return jsonify({'error': 'Unauthorized. You cannot delete a Super Admin account.'}), 403

    # Safety: Cannot delete yourself
    if user_to_delete.id == current_user.id:
        return jsonify({'error': 'Cannot delete your own account.'}), 400

    # Store info for audit log before deletion
    deleted_info = {
        'username': user_to_delete.username,
        'email': user_to_delete.email,
        'role': user_to_delete.role,
        'workspace_id': user_to_delete.workspace_id,
        'is_head_action': True # Mark as primary admin action
    }

    try:
        from app.models import (
            DeviceSession, Notification, Feedback, StudentProfile, 
            Assignment, GPARecord, TutorProfile, TutorSession, 
            GroupMessage, PasswordResetToken, Workspace,
            WorkspaceMembership, AssignmentGroupMember, AssignmentGroupMessage,
            File
        )
        # Import additional models from submodules that might block deletion
        from app.models.notifications import Notification as NotificationV2, NotificationPreference
        from app.models.academic import Submission
        from app.models.collaboration import Presence, FileVersion, CollaborationSession
        from app.models.document import Document, DocumentPermission, DocumentVersion
        from app.models.security import UserSession, TwoFactorAuth
        from app.models.chat_features import (
            MessageReaction, MessageReadReceipt, MessageDelivery,
            PinnedMessage, MessageForward, MessageEditHistory,
            ChannelPoll, PollVote, ChannelTopic, ScheduledMessage, ChannelMute
        )
        
        # 1. Unlink from Workspace if they are the head to avoid FK restriction
        Workspace.query.filter_by(admin_id=user_id).update({Workspace.admin_id: None})
        
        # 2. Advanced Manual cleanup of sub-module models (order matters for FKs)
        # Presence and Sessions
        Presence.query.filter_by(user_id=user_id).delete()
        UserSession.query.filter_by(user_id=user_id).delete()
        TwoFactorAuth.query.filter_by(user_id=user_id).delete()
        CollaborationSession.query.filter_by(user_id=user_id).delete()
        
        # Documents and Versions
        DocumentVersion.query.filter_by(created_by=user_id).delete()
        DocumentPermission.query.filter_by(user_id=user_id).delete()
        # Documents - if user owns them, they must go or be transferred. We delete them.
        Document.query.filter_by(owner_id=user_id).delete()
        
        # Notifications (both old and v2)
        NotificationV2.query.filter_by(user_id=user_id).delete()
        NotificationPreference.query.filter_by(user_id=user_id).delete()
        Notification.query.filter_by(user_id=user_id).delete()
        
        # Academic Submissions
        Submission.query.filter_by(student_id=user_id).delete()
        
        # Chat features and Messenger extensions
        MessageReaction.query.filter_by(user_id=user_id).delete()
        MessageReadReceipt.query.filter_by(user_id=user_id).delete()
        MessageDelivery.query.filter_by(user_id=user_id).delete()
        PinnedMessage.query.filter_by(pinned_by=user_id).delete()
        MessageForward.query.filter_by(forwarded_by=user_id).delete()
        MessageEditHistory.query.filter_by(edited_by=user_id).delete()
        ChannelPoll.query.filter_by(created_by=user_id).delete()
        PollVote.query.filter_by(user_id=user_id).delete()
        ChannelTopic.query.filter_by(created_by=user_id).delete()
        ScheduledMessage.query.filter_by(author_id=user_id).delete()
        ChannelMute.query.filter_by(user_id=user_id).delete()
        
        # File Related
        FileVersion.query.filter_by(uploaded_by=user_id).delete()
        File.query.filter_by(uploaded_by=user_id).delete()
        
        # Profile and Feedback
        DeviceSession.query.filter_by(user_id=user_id).delete()
        Feedback.query.filter((Feedback.student_id == user_id) | (Feedback.lecturer_id == user_id)).delete()
        StudentProfile.query.filter_by(user_id=user_id).delete()
        GPARecord.query.filter_by(user_id=user_id).delete()
        PasswordResetToken.query.filter_by(user_id=user_id).delete()
        
        # Tutor / Peer related
        TutorProfile.query.filter_by(user_id=user_id).delete()
        TutorSession.query.filter((TutorSession.tutor_id == user_id) | (TutorSession.student_id == user_id)).delete()
        GroupMessage.query.filter_by(user_id=user_id).delete()
        
        # Assignment related
        # Deleting assignments will cascade to groups if configured, but we clean groups members here too
        Assignment.query.filter_by(created_by=user_id).delete()
        AssignmentGroupMember.query.filter_by(user_id=user_id).delete()
        AssignmentGroupMessage.query.filter_by(user_id=user_id).delete()
        
        # Membership records
        WorkspaceMembership.query.filter_by(user_id=user_id).delete()
        
        # Finally delete the actual user record
        db.session.delete(user_to_delete)
        db.session.commit()
        
        # Log audit event
        log_audit_event(
            user_id=current_user_id,
            action='delete_user',
            resource_type='user',
            resource_id=user_id,
            details=deleted_info
        )
        
        return jsonify({'message': f'User {deleted_info["username"]} and all associated data deleted successfully'}), 200
    except Exception as e:
        db.session.rollback()
        import traceback
        print(f"User deletion error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Failed to delete user: {str(e)}'}), 500

