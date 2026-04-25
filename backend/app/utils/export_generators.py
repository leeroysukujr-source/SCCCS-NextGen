
import io
import os
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

def generate_pdf_report(data, title="System Report"):
    """
    Generate a high-fidelity, professional PDF report.
    Supports dynamic sections based on the gathered snapshot data.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                            rightMargin=50, leftMargin=50, 
                            topMargin=50, bottomMargin=50)
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Define Custom Styles
    styles.add(ParagraphStyle(
        name='MainTitle',
        fontSize=28,
        textColor=colors.HexColor('#1e293b'), # Slate-800
        alignment=0,
        fontName='Helvetica-Bold',
        spaceAfter=10
    ))
    
    styles.add(ParagraphStyle(
        name='SubTitle',
        fontSize=12,
        textColor=colors.HexColor('#64748b'), # Slate-500
        alignment=0,
        spaceAfter=30
    ))
    
    styles.add(ParagraphStyle(
        name='SectionHeader',
        fontSize=18,
        textColor=colors.HexColor('#4f46e5'), # Indigo-600
        fontName='Helvetica-Bold',
        spaceBefore=25,
        spaceAfter=15,
        borderPadding=10,
        borderWidth=0,
        borderColor=colors.HexColor('#e2e8f0'),
        backColor=colors.HexColor('#f8fafc')
    ))

    # --- Header ---
    if not isinstance(data, dict):
        data = {}
    metadata = data.get('metadata') if isinstance(data.get('metadata'), dict) else {}
    inst_name = metadata.get('institution', 'Institutional Report')
    elements.append(Paragraph(str(inst_name), styles['MainTitle']))
    elements.append(Paragraph(f"{title} | Generated on {datetime.now().strftime('%B %d, %Y')}", styles['SubTitle']))
    
    # --- Introduction ---
    elements.append(Paragraph("Executive Summary", styles['SectionHeader']))
    prepared_by = metadata.get('submitted_by', 'Administrator')
    elements.append(Paragraph(f"This report provides a comprehensive overview of the institutional performance and engagement metrics as requested. Submission prepared by {prepared_by}.", styles['Normal']))
    elements.append(Spacer(1, 15))

    # --- Metrics Logic ---
    metrics = data.get('metrics') if isinstance(data.get('metrics'), dict) else {}
    
    # 1. User Engagement Section
    if 'users' in metrics:
        elements.append(Paragraph("User Demographics & Enrollment", styles['SectionHeader']))
        u = metrics['users']
        user_table_data = [
            ['Metric Group', 'Count', 'Status'],
            ['Total Registered Users', str(u.get('total', 0)), 'Active'],
            ['Students', str(u.get('students', 0)), 'Enrolled'],
            ['Faculty/Teachers', str(u.get('teachers', 0)), 'Active'],
            ['Administrative Staff', str(u.get('admins', 0)), 'Active']
        ]
        
        t = Table(user_table_data, colWidths=[200, 100, 100])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4f46e5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor('#f8fafc')]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
        ]))
        elements.append(t)

    # 2. Activity Trends Section
    if 'activity' in metrics:
        elements.append(Paragraph("Platform Engagement Activity", styles['SectionHeader']))
        act = metrics['activity']
        elements.append(Paragraph(f"Analysis of daily interactions indicates a robust engagement level across the digital campus.", styles['Normal']))
        elements.append(Spacer(1, 10))
        
        act_data = [
            ['Activity Type', 'Volume', 'Context'],
            ['Daily Active Users (DAU)', str(act.get('daily_active_users', 0)), 'High Usage'],
            ['System Messages Transmitted', str(act.get('messages_sent', 0)), 'Active Channels'],
            ['Collaborative Meetings held', str(act.get('meetings_held', 0)), 'Live Sessions']
        ]
        t_act = Table(act_data, colWidths=[200, 100, 100])
        t_act.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
        ]))
        elements.append(t_act)

    # 3. Academic Performance
    if 'academic' in metrics:
        elements.append(Paragraph("Academic Operations & Quality", styles['SectionHeader']))
        acad = metrics['academic']
        elements.append(Paragraph(f"Operational data regarding courses and assessment throughput.", styles['Normal']))
        elements.append(Spacer(1, 10))
        
        acad_data = [
            ['Area', 'Metric', 'Performance'],
            ['Active Courses/Classes', str(acad.get('total_classes', 0)), 'Normal'],
            ['Assessment Submissions', str(acad.get('total_assignments', 0)), 'Steady'],
            ['Aggregate Performance Score', f"{acad.get('average_grade', 0)}%", 'Exceeding Target']
        ]
        t_acad = Table(acad_data, colWidths=[200, 100, 100])
        t_acad.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#059669')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
        ]))
        elements.append(t_acad)

    # --- Footer ---
    elements.append(Spacer(1, 50))
    elements.append(Paragraph("--- End of Official Institutional Report ---", styles['SubTitle']))
    elements.append(Paragraph("SCCCS NextGen Reporting Engine v2.0 | Secure Platform Verification", styles['SubTitle']))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_excel_report(data, title="Institutional Report"):
    """
    Generate a highly styled, multi-sheet Excel report.
    """
    wb = Workbook()
    
    # 1. Dashboard/Summary Sheet
    ws = wb.active
    ws.title = "Executive Summary"
    
    if not isinstance(data, dict):
        data = {}
    metadata = data.get('metadata') if isinstance(data.get('metadata'), dict) else {}
    metrics = data.get('metrics') if isinstance(data.get('metrics'), dict) else {}
    
    # Branding Header
    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=14)
    white_font = Font(color="FFFFFF")
    bold_font = Font(bold=True)
    
    ws['A1'] = metadata.get('institution', 'Institution')
    ws['A1'].font = header_font
    ws.merge_cells('A1:C1')
    
    ws.append([f"{title} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"])
    ws['A2'].font = Font(italic=True, color="64748B")
    ws.append([""]) # Spacer
    
    # User Demographics Section
    if 'users' in metrics:
        u = metrics['users']
        ws.append(["User Demographics & Enrollment"])
        ws[f"A{ws.max_row}"].font = Font(bold=True, size=12, color="4F46E5")
        
        headers = ["Metric Group", "Count", "Status"]
        ws.append(headers)
        for cell in ws[ws.max_row]:
            cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
            cell.font = white_font
            
        ws.append(["Total Registered Users", u.get('total', 0), "Active"])
        ws.append(["Students", u.get('students', 0), "Enrolled"])
        ws.append(["Faculty/Teachers", u.get('teachers', 0), "Active"])
        ws.append(["Administrative Staff", u.get('admins', 0), "Active"])
        ws.append([""]) # Spacer

    # Activity Trends Section
    if 'activity' in metrics:
        act = metrics['activity']
        ws.append(["Platform Engagement Activity"])
        ws[f"A{ws.max_row}"].font = Font(bold=True, size=12, color="0F172A")
        
        headers = ["Activity Type", "Volume", "Context"]
        ws.append(headers)
        for cell in ws[ws.max_row]:
            cell.fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
            cell.font = white_font
            
        ws.append(["Daily Active Users (DAU)", act.get('daily_active_users', 0), "High Usage"])
        ws.append(["System Messages Transmitted", act.get('messages_sent', 0), "Active Channels"])
        ws.append(["Collaborative Meetings held", act.get('meetings_held', 0), "Live Sessions"])
        ws.append([""]) # Spacer

    # Academic Section
    if 'academic' in metrics:
        acad = metrics['academic']
        ws.append(["Academic Operations & Quality"])
        ws[f"A{ws.max_row}"].font = Font(bold=True, size=12, color="059669")
        
        headers = ["Area", "Metric", "Performance"]
        ws.append(headers)
        for cell in ws[ws.max_row]:
            cell.fill = PatternFill(start_color="059669", end_color="059669", fill_type="solid")
            cell.font = white_font
            
        ws.append(["Active Courses/Classes", acad.get('total_classes', 0), "Normal"])
        ws.append(["Assessment Submissions", acad.get('total_assignments', 0), "Steady"])
        ws.append(["Aggregate Performance Score", f"{acad.get('average_grade', 0)}%", "Exceeding Target"])

    # Auto-width
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except: pass
        ws.column_dimensions[column_letter].width = min(max_length + 2, 40)
            
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_word_report(data, title="Institutional Report"):
    """
    Generate a professional, captivating Word document.
    """
    doc = Document()
    
    # Configure document-wide styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    if not isinstance(data, dict):
        data = {}
    metadata = data.get('metadata') if isinstance(data.get('metadata'), dict) else {}
    metrics = data.get('metrics') if isinstance(data.get('metrics'), dict) else {}

    # 1. Header
    section = doc.sections[0]
    header = section.header
    htab = header.paragraphs[0]
    htab.text = f"Official Institutional Record | {datetime.now().strftime('%Y-%m-%d')}"
    htab.style = doc.styles['Caption']

    # 2. Main Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(metadata.get('institution', 'Institutional Report'))
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    title_p.alignment = 0 # Left

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"{title}\nGenerated via SCCCS NextGen Reporting Engine")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph("-" * 80)

    # 3. Content Sections
    def add_section_table(doc, title, data_list, theme_color):
        doc.add_heading(title, level=1)
        
        table = doc.add_table(rows=len(data_list), cols=len(data_list[0]))
        table.style = 'Table Grid'
        
        # Style Header Row
        for i, text in enumerate(data_list[0]):
            cell = table.rows[0].cells[i]
            cell.text = text
            shading_elm = qn('w:shd')
            shading_elm.set(qn('w:fill'), theme_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            p = cell.paragraphs[0]
            run = p.add_run()
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True

        # Data Rows
        for i, row_data in enumerate(data_list[1:], start=1):
            for j, val in enumerate(row_data):
                table.rows[i].cells[j].text = str(val)

    # User Demographics
    if 'users' in metrics:
        u = metrics['users']
        user_rows = [
            ['Metric Group', 'Count', 'Status'],
            ['Total Registered Users', u.get('total', 0), 'Active'],
            ['Students', u.get('students', 0), 'Enrolled'],
            ['Faculty/Teachers', u.get('teachers', 0), 'Active'],
            ['Administrative Staff', u.get('admins', 0), 'Active']
        ]
        add_section_table(doc, "User Demographics & Enrollment", user_rows, "4F46E5") # Indigo

    doc.add_paragraph("\n")

    # Activity
    if 'activity' in metrics:
        act = metrics['activity']
        act_rows = [
            ['Activity Type', 'Volume', 'Context'],
            ['Daily Active Users (DAU)', act.get('daily_active_users', 0), 'High Usage'],
            ['System Messages Transmitted', act.get('messages_sent', 0), 'Active'],
            ['Collaborative Meetings held', act.get('meetings_held', 0), 'Live']
        ]
        add_section_table(doc, "Platform Engagement Activity", act_rows, "0F172A") # Dark Slate

    doc.add_paragraph("\n")

    # Academic
    if 'academic' in metrics:
        acad = metrics['academic']
        acad_rows = [
            ['Area', 'Metric', 'Performance'],
            ['Active Courses/Classes', acad.get('total_classes', 0), 'Normal'],
            ['Assessment Submissions', acad.get('total_assignments', 0), 'Steady'],
            ['Aggregate Performance Score', f"{acad.get('average_grade', 0)}%", 'High']
        ]
        add_section_table(doc, "Academic Operations & Quality", acad_rows, "059669") # Emerald

    # 4. Footer
    doc.add_paragraph("\n\n" + "-" * 80)
    footer_p = doc.add_paragraph("--- End of Official Institutional Report ---")
    footer_p.alignment = 1
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
